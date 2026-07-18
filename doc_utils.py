from docx import Document
from io import BytesIO


def create_docx(email_content):
    """
    Creates a DOCX file from the generated email
    and returns it as bytes.
    """

    doc = Document()

    doc.add_heading("Generated Email", level=1)

    doc.add_paragraph(email_content)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer